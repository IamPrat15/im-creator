"""
IM Creator Python Backend - FastAPI Application
Version: 7.2.0

Main entry point with all API endpoints.
"""

import os
import tempfile
import uuid
import base64

from usage_tracker import get_tracker
from state_manager import PresentationStateManager
from fastapi.responses import Response

from datetime import datetime
from pathlib import Path
from typing import Optional, Dict, Any

from fastapi import FastAPI, HTTPException, Response
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse
from pydantic import BaseModel
from dotenv import load_dotenv

'''from models import (
    VERSION, PROFESSIONAL_TEMPLATES, INDUSTRY_DATA, DOCUMENT_CONFIGS,
    GeneratePPTXRequest, GeneratePPTXRequestWrapper
)'''
from models import (
    VERSION, PROFESSIONAL_TEMPLATES, INDUSTRY_DATA, DOCUMENT_CONFIGS
)

from pptx_generator import generate_presentation
from ai_layout_engine import get_usage_stats, reset_usage_stats

# Load environment variables
load_dotenv()

# ============================================================================
# APP INITIALIZATION
# ============================================================================

app = FastAPI(
    title="IM Creator API",
    description="AI-Powered Information Memorandum Creator - Python Backend",
    version=VERSION["string"]
)

# CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Configure appropriately for production
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Temp directory for generated files
TEMP_DIR = Path(tempfile.gettempdir()) / "im_creator"
TEMP_DIR.mkdir(exist_ok=True)

# ============================================================================
# HEALTH & INFO ENDPOINTS
# ============================================================================

@app.get("/api/health")
async def health_check():
    """Health check endpoint"""
    return {
        "status": "ok",
        "version": VERSION["string"],
        "version_full": VERSION["full"],
        "build_date": VERSION["build_date"],
        "backend": "Python/FastAPI",
        "pptx_library": "python-pptx",
        "features": [
            "AI-Powered Layout Engine",
            "Native Pie/Donut/Bar Charts (no fallback)",
            "Progress Bars and Stacked Bars",
            "50 Professional Templates",
            "6 Industry Verticals",
            "3 Document Types",
            "Dynamic Font Adjustment"
        ]
    }

@app.get("/api/version")
async def get_version():
    """Get version information"""
    return {
        "current": VERSION["string"],
        "full": VERSION["full"],
        "build_date": VERSION["build_date"],
        "history": VERSION["history"]
    }

# ============================================================================
# TEMPLATES & CONFIG ENDPOINTS
# ============================================================================

@app.get("/api/templates")
async def get_templates():
    """Get all available templates"""
    return PROFESSIONAL_TEMPLATES

@app.get("/api/industries")
async def get_industries():
    """Get all industry/vertical data"""
    return list(INDUSTRY_DATA.values())

@app.get("/api/document-types")
async def get_document_types():
    """Get document type configurations"""
    return DOCUMENT_CONFIGS

# ============================================================================
# USAGE TRACKING ENDPOINTS
# ============================================================================

@app.get("/api/usage")
async def get_usage():
    """Get AI API usage statistics"""
    return get_usage_stats()

@app.post("/api/usage/reset")
async def reset_usage():
    """Reset usage statistics"""
    reset_usage_stats()
    return {"success": True, "message": "Usage statistics reset"}

@app.get("/api/usage/export")
async def export_usage():
    """Export usage to CSV"""
    stats = get_usage_stats()
    
    headers = ["Timestamp", "Model", "Input Tokens", "Output Tokens", "Cost (USD)", "Purpose"]
    rows = []
    for call in stats.get("recent_calls", []):
        rows.append([
            call.get("timestamp", ""),
            call.get("model", ""),
            str(call.get("input_tokens", 0)),
            str(call.get("output_tokens", 0)),
            call.get("cost_usd", "0"),
            call.get("purpose", "")
        ])
    
    csv_content = ",".join(headers) + "\n"
    for row in rows:
        csv_content += ",".join([f'"{v}"' for v in row]) + "\n"
    
    return Response(
        content=csv_content,
        media_type="text/csv",
        headers={"Content-Disposition": "attachment; filename=usage_report.csv"}
    )

# ============================================================================
# PPTX GENERATION ENDPOINT
# ============================================================================

class GenerateRequest(BaseModel):
    data: Dict[str, Any]
    theme: Optional[str] = "modern-blue"

'''@app.post("/api/generate-pptx")
async def generate_pptx(request: GenerateRequest):
    """Generate PowerPoint presentation"""
    try:
        data = request.data
        theme = request.theme or "modern-blue"
        
        print(f"\n{'='*50}")
        print(f"Generating PPTX with Python Backend v{VERSION['string']}")
        print(f"Theme: {theme}")
        print(f"Document Type: {data.get('documentType') or data.get('document_type') or 'management-presentation'}")
        print(f"Company: {data.get('companyName') or data.get('company_name') or 'Unknown'}")
        print(f"{'='*50}")
        
        # Generate presentation
        prs = await generate_presentation(data, theme)
        
        # Generate filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        codename = (data.get("projectCodename") or data.get("project_codename") or "Project").replace(" ", "_")
        codename = "".join(c for c in codename if c.isalnum() or c == "_")
        filename = f"{codename}_{timestamp}.pptx"
        filepath = TEMP_DIR / filename
        
        # Save presentation
        prs.save(str(filepath))
        
        print(f"Generated: {filename}")
        print(f"File size: {filepath.stat().st_size / 1024:.1f} KB")
        
        # Return file
        return FileResponse(
            path=str(filepath),
            filename=filename,
            media_type="application/vnd.openxmlformats-officedocument.presentationml.presentation",
            background=None  # File will be deleted after response
        )
        
    except Exception as e:
        print(f"PPTX generation error: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))
'''
@app.post("/api/generate-pptx")
async def generate_pptx(request: GenerateRequest):
    """Generate PowerPoint presentation"""
    try:
        data = request.data
        theme = request.theme or "modern-blue"
        
        print(f"\n{'='*50}")
        print(f"Generating PPTX with Python Backend v{VERSION['string']}")
        print(f"Theme: {theme}")
        print(f"Document Type: {data.get('documentType') or data.get('document_type') or 'management-presentation'}")
        print(f"Company: {data.get('companyName') or data.get('company_name') or 'Unknown'}")
        print(f"{'='*50}")
        

        # ========================================
        # REQUIREMENT #7: CONDITIONAL VALIDATION
        # ========================================
        buyer_types = data.get("targetBuyerType", [])
        if isinstance(buyer_types, str):
            buyer_types = [buyer_types]
        
        validation_errors = []
        
        # Rule 1: Strategic buyers need synergies
        if "strategic" in buyer_types and not data.get("synergiesStrategic") and not data.get("synergies_strategic"):
            validation_errors.append("Strategic synergies are required when targeting strategic buyers")
        
        # Rule 2: Financial buyers need financial synergies
        if "financial" in buyer_types and not data.get("synergiesFinancial") and not data.get("synergies_financial"):
            validation_errors.append("Financial synergies are required when targeting financial investors")
        
        # Rule 3: CIM requires additional details
        doc_type = (data.get("documentType") or data.get("document_type") or "").lower()
        if doc_type == "cim":
            if not data.get("leadershipTeam") and not data.get("leadership_team"):
                validation_errors.append("Leadership team details are required for CIM documents")
            if not data.get("competitiveAdvantages") and not data.get("competitive_advantages"):
                validation_errors.append("Competitive advantages are required for CIM documents")
            if not data.get("growthDrivers") and not data.get("growth_drivers"):
                validation_errors.append("Growth drivers are required for CIM documents")
        
        # Rule 4: If revenue provided, margins required
        if (data.get("revenueFY25") or data.get("revenue_fy25")):
            if not data.get("ebitdaMarginFY25") and not data.get("ebitda_margin_fy25"):
                validation_errors.append("EBITDA margin is required when FY25 revenue is provided")
        
        # Rule 5: If including additional case studies, must have at least 3 total
        if data.get("includeAdditionalCaseStudies") or data.get("include_additional_case_studies"):
            case_studies = data.get("caseStudies") or []
            if len(case_studies) < 3:
                validation_errors.append("At least 3 case studies required to include additional case studies in appendix")
        
        # Return validation errors if any
        if validation_errors:
            print(f"Validation failed: {validation_errors}")
            return JSONResponse(
                status_code=400,
                content={
                    "success": False,
                    "error": "Validation failed",
                    "validation_errors": validation_errors,
                    "message": "Please fill in all required fields based on your selections"
                }
            )

        # Generate presentation
        prs = generate_presentation(data, theme)
        
        # Generate filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        codename = (data.get("projectCodename") or data.get("project_codename") or "Project").replace(" ", "_")
        codename = "".join(c for c in codename if c.isalnum() or c == "_")
        filename = f"{codename}_{timestamp}.pptx"
        filepath = TEMP_DIR / filename
        
        # Save presentation
        prs.save(str(filepath))
        
        print(f"Generated: {filename}")
        print(f"File size: {filepath.stat().st_size / 1024:.1f} KB")
        
        # Read file and convert to base64
        import base64
        with open(filepath, "rb") as f:
            file_bytes = f.read()
            file_base64 = base64.b64encode(file_bytes).decode('utf-8')
        
        # Count slides
        slide_count = len(prs.slides)
        
        # Clean up temp file
        try:
            filepath.unlink()
        except:
            pass
        
        # Return JSON response with base64 data
        return JSONResponse({
            "success": True,
            "fileData": file_base64,
            "filename": filename,
            "mimeType": "application/vnd.openxmlformats-officedocument.presentationml.presentation",
            "slideCount": slide_count,
            "message": f"Generated {slide_count} slides successfully"
        })
        
    except Exception as e:
        print(f"PPTX generation error: {e}")
        import traceback
        traceback.print_exc()
        return JSONResponse(
            status_code=500,
            content={
                "success": False,
                "error": str(e),
                "message": "Failed to generate presentation"
            }
        )
# ============================================================================
# REQUIREMENT #12: WORD Q&A EXPORT ENDPOINT
# ============================================================================

class ExportQARequest(BaseModel):
    data: Dict[str, Any]
    questionnaire: Optional[Dict[str, Any]] = None

@app.post("/api/export-qa-word")
async def export_qa_word(request: ExportQARequest):
    """
    Export Questions and Answers to Word format
    Implements Requirement #12
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor as DocxRGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # Create document
        doc = Document()
        
        # Title
        title = doc.add_heading('Information Memorandum - Questions & Answers', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadata
        company = request.data.get("companyName") or "Company"
        codename = request.data.get("projectCodename") or "Project"
        
        doc.add_paragraph(f"Company: {company}")
        doc.add_paragraph(f"Project: {codename}")
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
        doc.add_paragraph()  # Blank line
        
        # Add sections
        sections = [
            ("Project Setup", [
                ("Project Codename", request.data.get("projectCodename")),
                ("Company Name", request.data.get("companyName")),
                ("Document Type", request.data.get("documentType")),
                ("Advisor", request.data.get("advisor")),
                ("Presentation Date", request.data.get("presentationDate")),
            ]),
            ("Company Overview", [
                ("Founded Year", request.data.get("foundedYear")),
                ("Headquarters", request.data.get("headquarters")),
                ("Company Description", request.data.get("companyDescription")),
                ("Full-Time Employees", request.data.get("employeeCountFT")),
                ("Investment Highlights", request.data.get("investmentHighlights")),
            ]),
            ("Leadership", [
                ("Founder Name", request.data.get("founderName")),
                ("Founder Title", request.data.get("founderTitle")),
                ("Years of Experience", request.data.get("founderExperience")),
                ("Education", request.data.get("founderEducation")),
                ("Leadership Team", request.data.get("leadershipTeam")),
            ]),
            ("Services & Products", [
                ("Service Lines", request.data.get("serviceLines")),
                ("Products", request.data.get("products")),
                ("Tech Partnerships", request.data.get("techPartnerships")),
            ]),
            ("Clients", [
                ("Primary Vertical", request.data.get("primaryVertical")),
                ("Top Clients", request.data.get("topClients")),
                ("Top 10 Concentration", request.data.get("top10Concentration")),
                ("Net Revenue Retention", request.data.get("netRetention")),
            ]),
            ("Financials", [
                ("Currency", request.data.get("currency")),
                ("Revenue FY24", request.data.get("revenueFY24")),
                ("Revenue FY25", request.data.get("revenueFY25")),
                ("Revenue FY26P", request.data.get("revenueFY26P")),
                ("EBITDA Margin FY25", request.data.get("ebitdaMarginFY25")),
                ("Gross Margin", request.data.get("grossMargin")),
            ]),
            ("Growth Strategy", [
                ("Growth Drivers", request.data.get("growthDrivers")),
                ("Competitive Advantages", request.data.get("competitiveAdvantages")),
                ("Short-Term Goals", request.data.get("shortTermGoals")),
                ("Medium-Term Goals", request.data.get("mediumTermGoals")),
            ]),
        ]
        
        for section_name, questions in sections:
            doc.add_heading(section_name, level=1)
            
            for question, answer in questions:
                if answer:
                    # Question
                    q_para = doc.add_paragraph()
                    q_run = q_para.add_run(f"Q: {question}")
                    q_run.bold = True
                    q_run.font.size = Pt(11)
                    
                    # Answer
                    a_para = doc.add_paragraph()
                    a_run = a_para.add_run(f"A: {answer}")
                    a_run.font.size = Pt(10)
                    
                    doc.add_paragraph()  # Blank line
        
        # Save to temp file
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{codename}_QA_{timestamp}.docx"
        filepath = TEMP_DIR / filename
        doc.save(str(filepath))
        
        # Return file
        return FileResponse(
            path=str(filepath),
            filename=filename,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
    except Exception as e:
        print(f"Word export error: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))


# ============================================================================
# REQUIREMENT #13: PDF EXPORT ENDPOINT
# ============================================================================

@app.post("/api/export-pdf")
async def export_pdf(request: Dict[str, Any]):
    """
    Export presentation to PDF format
    Implements Requirement #13 (partial - basic implementation)
    """
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        
        # Get data
        data = request.get("data", {})
        
        # Generate filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        codename = data.get("projectCodename") or "Project"
        filename = f"{codename}_{timestamp}.pdf"
        filepath = TEMP_DIR / filename
        
        # Create PDF
        doc = SimpleDocTemplate(str(filepath), pagesize=letter)
        story = []
        styles = getSampleStyleSheet()
        
        # Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#2B579A'),
            spaceAfter=30,
            alignment=1  # Center
        )
        
        company = data.get("companyName") or "Company Name"
        story.append(Paragraph(company, title_style))
        story.append(Spacer(1, 0.5*inch))
        
        # Add sections
        sections = [
            ("Company Overview", data.get("companyDescription") or ""),
            ("Investment Highlights", data.get("investmentHighlights") or ""),
            ("Services", data.get("serviceLines") or ""),
            ("Growth Strategy", data.get("growthDrivers") or ""),
        ]
        
        for section_title, content in sections:
            if content:
                story.append(Paragraph(section_title, styles['Heading2']))
                story.append(Spacer(1, 0.2*inch))
                story.append(Paragraph(content, styles['Normal']))
                story.append(Spacer(1, 0.3*inch))
        
        # Build PDF
        doc.build(story)
        
        # Return file
        return FileResponse(
            path=str(filepath),
            filename=filename,
            media_type="application/pdf"
        )
        
    except Exception as e:
        print(f"PDF export error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ============================================================================
# CHAT ENDPOINT (for AI assistant)
# ============================================================================

class ChatRequest(BaseModel):
    message: str
    conversation_history: Optional[list] = []
    context: Optional[Dict[str, Any]] = {}

@app.post("/api/chat")
async def chat(request: ChatRequest):
    """Chat with AI assistant"""
    try:
        from anthropic import Anthropic
        
        api_key = os.environ.get("ANTHROPIC_API_KEY")
        if not api_key:
            raise HTTPException(status_code=500, detail="Anthropic API key not configured")
        
        client = Anthropic(api_key=api_key)
        
        system_prompt = f"""You are an expert M&A advisor helping create professional Information Memorandum presentations. 
        You help users fill out company information for CIM, Management Presentations, and Teasers.
        Be concise, professional, and helpful. Ask clarifying questions when needed.
        Current form context: {request.context}"""
        
        messages = [
            *[{"role": msg.get("role"), "content": msg.get("content")} for msg in request.conversation_history],
            {"role": "user", "content": request.message}
        ]
        
        response = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=1024,
            system=system_prompt,
            messages=messages
        )
        
        return {
            "response": response.content[0].text,
            "usage": {
                "input_tokens": response.usage.input_tokens,
                "output_tokens": response.usage.output_tokens
            }
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# ============================================================================
# DRAFT MANAGEMENT
# ============================================================================

DRAFTS_DIR = TEMP_DIR / "drafts"
DRAFTS_DIR.mkdir(exist_ok=True)

class DraftRequest(BaseModel):
    data: Dict[str, Any]
    project_id: str

@app.post("/api/drafts")
async def save_draft(request: DraftRequest):
    """Save draft"""
    try:
        import json
        draft_path = DRAFTS_DIR / f"{request.project_id}.json"
        with open(draft_path, "w") as f:
            json.dump(request.data, f, indent=2)
        return {"success": True, "project_id": request.project_id}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/drafts/{project_id}")
async def get_draft(project_id: str):
    """Get draft by project ID"""
    try:
        import json
        draft_path = DRAFTS_DIR / f"{project_id}.json"
        if not draft_path.exists():
            raise HTTPException(status_code=404, detail="Draft not found")
        with open(draft_path, "r") as f:
            data = json.load(f)
        return {"success": True, "data": data}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# ============================================================================
# STARTUP
# ============================================================================

@app.on_event("startup")
async def startup_event():
    """Startup event"""
    print(f"\n{'='*60}")
    print(f"IM Creator Python Backend {VERSION['full']}")
    print(f"{'='*60}")
    print(f"Using python-pptx for native chart generation")
    print(f"\nFeatures:")
    print(f"  ✓ Native Pie/Donut/Bar charts (no fallback)")
    print(f"  ✓ AI-powered layout recommendations")
    print(f"  ✓ 50 professional templates")
    print(f"  ✓ 6 industry verticals")
    print(f"  ✓ 3 document types")
    print(f"{'='*60}\n")

@app.get("/api/usage")
async def get_usage():
    """Get Anthropic API usage statistics"""
    tracker = get_tracker()
    return tracker.get_stats()

@app.get("/api/usage/export")
async def export_usage():
    """Export usage data as CSV"""
    tracker = get_tracker()
    csv_data = tracker.export_csv()
    return Response(content=csv_data, media_type="text/csv",
                   headers={"Content-Disposition": "attachment; filename=usage.csv"})

@app.post("/api/usage/reset")
async def reset_usage():
    """Reset usage statistics"""
    tracker = get_tracker()
    tracker.reset()
    return {"success": True, "message": "Usage statistics reset"}

@app.post("/api/update-presentation")
async def update_presentation(request: dict):
    """Dynamic slide updates - Requirement #8"""
    manager = PresentationStateManager()
    changed = manager.detect_changes(request["old_data"], request["new_data"])
    if not changed:
        return {"message": "No changes detected"}
    affected = manager.get_affected_slides(changed)
    return {"slides_updated": len(affected), "affected_slides": affected}

# ============================================================================
# MAIN
# ============================================================================

if __name__ == "__main__":
    import uvicorn
    port = int(os.environ.get("PORT", 8000))
    uvicorn.run(app, host="0.0.0.0", port=port)
